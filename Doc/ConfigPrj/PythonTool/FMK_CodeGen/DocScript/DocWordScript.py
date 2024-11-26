"""
#  @file        main.py
#  @brief       Template_BriefDescription.
#  @details     TemplateDetailsDescription.\n
#
#  @author      mba
#  @date        jj/mm/yyyy
#  @version     1.0
"""
#------------------------------------------------------------------------------
#                                       IMPORT
#------------------------------------------------------------------------------
import os
import re
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
#------------------------------------------------------------------------------
#                                       CONSTANT
#------------------------------------------------------------------------------
# CAUTION : Automatic generated code section: Start #

# CAUTION : Automatic generated code section: End #
#------------------------------------------------------------------------------
#                                       CLASS
#------------------------------------------------------------------------------

# Fonction pour ajouter un style personnalisé pour le code
def add_code_style(doc):
    styles = doc.styles
    try:
        code_style = styles['Code']
    except KeyError:
        # Créer un style "Code" s'il n'existe pas
        code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        font = code_style.font
        font.name = 'Courier New'
        font.size = Pt(10.5)  # Taille adaptée pour du code
        code_style.paragraph_format.left_indent = Pt(10)
        code_style.paragraph_format.right_indent = Pt(10)
    return code_style

# Fonction pour ajouter un style de titre
def add_heading_style(doc):
    styles = doc.styles
    heading_style = styles['Heading 1']
    heading_style.font.size = Pt(14)
    heading_style.font.bold = True
    return heading_style

# Fonction pour extraire les informations d'une fonction
def extract_function_info(function_line):
    function_info = {
        "brief": "",
        "note": "",
        "params": [],
        "retval": []
    }

    # Regex pour extraire les éléments @brief, @note, @param et @retval
    brief_match = re.search(r'@brief\s+(.+)', function_line)
    note_match = re.search(r'@note\s+(.+)', function_line)
    param_matches = re.findall(r'@param\[[^\]]*\]\s+:\s+(.+)', function_line)
    retval_matches = re.findall(r'@retval\s+(.+)', function_line)

    if brief_match:
        function_info["brief"] = brief_match.group(1)
    if note_match:
        function_info["note"] = note_match.group(1)
    if param_matches:
        function_info["params"] = param_matches
    if retval_matches:
        function_info["retval"] = retval_matches

    return function_info

# Fonction pour scanner les fichiers .h et extraire les informations des fonctions
def scan_headers_and_generate_doc(header_dir, output_file):
    doc = Document()
    add_code_style(doc)
    
    # Titre de niveau 1 pour chaque fichier
    heading_style = add_heading_style(doc)

    for filename in os.listdir(header_dir):
        if filename.endswith(".h"):
            file_path = os.path.join(header_dir, filename)
            doc.add_paragraph(filename, style=heading_style)  # Titre 1 - Nom du fichier

            with open(file_path, 'r') as file:
                lines = file.readlines()
                functions = []
                function_info = {}

                # Parcours des lignes du fichier pour extraire les prototypes de fonctions
                for line in lines:
                    # Trouver les prototypes de fonctions
                    if line.strip().startswith("t_eReturnCode") or line.strip().startswith("void"):
                        functions.append(line.strip())
                
                # Traitement de chaque fonction
                for func_line in functions:
                    function_info = extract_function_info(func_line)
                    
                    # Titre 2 pour le nom de la fonction
                    doc.add_paragraph(func_line.split('(')[0], style='Heading 2')  # Nom de la fonction
                    
                    # Description de la fonction (brief et note)
                    doc.add_paragraph("Description", style='Heading 3')
                    doc.add_paragraph(f"Brief: {function_info['brief']}")
                    doc.add_paragraph(f"Note: {function_info['note']}")
                    
                    # Paramètres de la fonction
                    doc.add_paragraph("Paramètres", style='Heading 3')
                    for param in function_info["params"]:
                        doc.add_paragraph(param)
                    
                    # Valeurs de retour de la fonction
                    doc.add_paragraph("Retour", style='Heading 3')
                    for retval in function_info["retval"]:
                        doc.add_paragraph(retval)
    
    # Sauvegarder le document
    doc.save(output_file)


#------------------------------------------------------------------------------
#                             FUNCTION IMPLMENTATION
#------------------------------------------------------------------------------
def main()-> None:
    # Chemin du répertoire contenant les fichiers .h
    header_dir = "src\\1_FMK\FMK_HAL\FMK_CPU"

    # Nom du fichier Word généré
    output_file = "Doc\InfoPrj\Documentation_des_fichiers_h.docx"

    # Exécution du script
    scan_headers_and_generate_doc(header_dir, output_file)
    return
#------------------------------------------------------------------------------
#			                MAIN
#------------------------------------------------------------------------------
if (__name__ == '__main__'):
    main()

#------------------------------------------------------------------------------
#		                    END OF FILE
#------------------------------------------------------------------------------
#--------------------------
# Function_name
#--------------------------

"""
    @brief
    @details

    @params[in]
    @params[out]
    @retval
"""

